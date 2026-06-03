"""集成测试：验证完整流程"""
import pytest
from pathlib import Path
import tempfile
import json

from docx import Document

from config import load_config
from workflow import build_workflow, init_state
from parser import parse_inner_docx, parse_outer_docx


def create_test_data():
    """创建测试数据"""
    # 创建临时目录
    outer_dir = Path(tempfile.mkdtemp()) / "outer"
    inner_dir = Path(tempfile.mkdtemp()) / "inner"
    outer_dir.mkdir(parents=True)
    inner_dir.mkdir(parents=True)

    # 创建外规文档
    doc = Document()
    doc.add_paragraph("《测试外规》")
    doc.add_paragraph("第一条 外规条款内容关于贷款审批。")
    doc.add_paragraph("第二条 外规条款内容关于风险管理。")
    doc.save(str(outer_dir / "《测试外规》.docx"))

    # 创建内规文档
    doc = Document()
    doc.add_paragraph("第一条 内规条款内容关于贷款审批。")
    doc.add_paragraph("第二条 内规条款内容关于风险控制。")
    doc.save(str(inner_dir / "《测试内规》.docx"))

    return outer_dir, inner_dir


@pytest.mark.skip(reason="需要 API Key")
def test_full_workflow():
    """测试完整工作流（需要 API Key）"""
    outer_dir, inner_dir = create_test_data()
    output_dir = Path(tempfile.mkdtemp())

    config = {
        "data": {
            "clauses_json_path": "",
            "outer_dir": str(outer_dir),
            "inner_dir": str(inner_dir),
            "output_dir": str(output_dir),
        },
        "filter": {
            "hard_filter": {},
            "soft_filter": {"min_score": 0.0, "score_weights": {}},
        },
        "llm": {
            "judgment_model": "qwen-plus",
            "review_model": "qwen-plus",
        },
        "review": {
            "pass_threshold": 70,
            "max_retries": 2,
        },
        "split": {
            "val_clause_ratio": 0.5,
            "random_seed": 42,
        },
    }

    workflow = build_workflow()
    app = workflow.compile()
    initial_state = init_state(config)

    result = app.invoke(initial_state)

    # 验证输出
    assert len(result.get("qa_pairs", [])) > 0 or len(result.get("failed_pairs", [])) > 0
    assert (output_dir / "train.jsonl").exists() or (output_dir / "failed_pairs.jsonl").exists()


def test_parser_integration():
    """测试解析器集成"""
    outer_dir, inner_dir = create_test_data()

    # 测试外规解析
    outer_files = list(outer_dir.glob("*.docx"))
    outer_result = parse_outer_docx(outer_files[0])
    assert len(outer_result["clauses"]) == 2

    # 测试内规解析
    inner_files = list(inner_dir.glob("*.docx"))
    inner_clauses = parse_inner_docx(inner_files[0])
    assert len(inner_clauses) == 2


def test_workflow_state_initialization():
    """测试工作流状态初始化"""
    config = {"data": {"output_dir": "output"}}
    state = init_state(config)

    assert state["outer_clauses"] == []
    assert state["inner_clauses"] == []
    assert state["retry_count"] == 0


def test_filter_integration():
    """测试筛选模块集成"""
    from filter import compute_similarity, check_hard_filter, compute_weighted_score

    outer = {
        "subjects": ["商业银行"],
        "objects": ["贷款"],
        "biz_dims": ["信贷管理"],
        "action": "不得",
    }
    inner = {
        "subjects": ["商业银行"],
        "objects": ["贷款"],
        "biz_dims": ["信贷管理"],
        "action": "应当",
    }

    similarity = compute_similarity(outer, inner)
    assert similarity["subject_match"] == 1
    assert similarity["biz_dim_match"] == 1

    # 硬筛选
    hard_config = {"biz_dim_match": ">= 1"}
    assert check_hard_filter(similarity, hard_config) is True

    # 软筛选
    weights = {
        "biz_dim_match": 0.30,
        "subject_match": 0.25,
        "object_match": 0.25,
        "action_match": 0.10,
        "text_similarity": 0.10,
    }
    score = compute_weighted_score(similarity, weights)
    assert score > 0


def test_sft_converter_integration():
    """测试 SFT 转换器集成"""
    from utils.sft_converter import convert_to_sft_format

    qa_pair = {
        "outer_clause": {
            "doc_title": "商业银行法",
            "article_no": "第四十条",
            "clause_text": "商业银行不得向关系人发放信用贷款",
            "subjects": ["商业银行"],
            "objects": ["信用贷款"],
            "biz_dims": ["信贷管理"],
            "action": "不得",
        },
        "inner_clause": {
            "doc_title": "信贷制度",
            "article_no": "第十五条",
            "clause_text": "本行向关联方发放贷款应当审查",
            "subjects": ["商业银行"],
            "objects": ["贷款"],
            "biz_dims": ["信贷管理"],
            "action": "应当",
        },
        "label": "部分冲突",
        "confidence": 0.85,
        "reasoning": "行为模态不一致",
        "conflict_details": {"conflict_type": "行为模态冲突", "description": "外规禁止，内规要求审查"},
    }

    sft = convert_to_sft_format(qa_pair)

    assert "instruction" in sft
    assert "input" in sft
    assert "output" in sft

    output = json.loads(sft["output"])
    assert output["label"] == "部分冲突"
    assert output["conflict_type"] == "行为模态冲突"


def test_config_loading():
    """测试配置加载集成"""
    config = load_config("config.yaml")

    assert config.data.outer_dir == "data/outer"
    assert config.data.inner_dir == "data/inner"
    assert config.llm.judgment_model == "qwen-plus"
    assert config.review.pass_threshold == 70


def test_workflow_graph_building():
    """测试工作流图构建"""
    workflow = build_workflow()

    # 验证工作流可编译
    app = workflow.compile()
    assert app is not None


def test_end_to_end_sft_conversion():
    """测试端到端 SFT 转换"""
    from utils.sft_converter import convert_dataset_to_sft

    # 创建模拟数据集
    dataset = [
        {
            "outer_clause": {
                "doc_title": "商业银行法",
                "article_no": "第四十条",
                "clause_text": "商业银行不得向关系人发放信用贷款",
                "subjects": ["商业银行"],
                "objects": ["信用贷款"],
                "biz_dims": ["信贷管理"],
                "action": "不得",
            },
            "inner_clause": {
                "doc_title": "信贷管理制度",
                "article_no": "第十五条",
                "clause_text": "本行禁止向关系人发放信用贷款",
                "subjects": ["本行"],
                "objects": ["信用贷款"],
                "biz_dims": ["信贷管理"],
                "action": "禁止",
            },
            "label": "完全符合",
            "confidence": 0.95,
            "reasoning": "内规完全落实外规要求，行为模态一致",
        },
        {
            "outer_clause": {
                "doc_title": "银行业监督管理法",
                "article_no": "第二十一条",
                "clause_text": "银行业金融机构应当按规定披露信息",
                "subjects": ["银行业金融机构"],
                "objects": ["信息"],
                "biz_dims": ["信息披露"],
                "action": "应当",
            },
            "inner_clause": {
                "doc_title": "信息披露制度",
                "article_no": "第八条",
                "clause_text": "本行每季度披露财务报告和风险信息",
                "subjects": ["本行"],
                "objects": ["财务报告", "风险信息"],
                "biz_dims": ["信息披露"],
                "action": "披露",
            },
            "label": "基本一致",
            "confidence": 0.85,
            "reasoning": "内规落实外规要求，但披露频率有具体规定",
        },
    ]

    sft_dataset = convert_dataset_to_sft(dataset)

    assert len(sft_dataset) == 2

    # 验证每条数据的格式
    for sft in sft_dataset:
        assert "instruction" in sft
        assert "input" in sft
        assert "output" in sft

        # 验证 input 是有效的 JSON
        input_data = json.loads(sft["input"])
        assert "outer_clause" in input_data
        assert "inner_clause" in input_data

        # 验证 output 是有效的 JSON
        output_data = json.loads(sft["output"])
        assert "label" in output_data
        assert "confidence" in output_data


def test_similarity_filter_pipeline():
    """测试相似度计算到筛选的完整流程"""
    from filter import compute_similarity, check_hard_filter, compute_weighted_score, filter_candidates

    # 创建候选对
    outer_clauses = [
        {
            "uid": "outer_1",
            "subjects": ["商业银行"],
            "objects": ["贷款"],
            "biz_dims": ["信贷管理"],
            "action": "不得",
            "clause_text": "商业银行不得向关系人发放信用贷款",
        },
        {
            "uid": "outer_2",
            "subjects": ["银行业金融机构"],
            "objects": ["信息"],
            "biz_dims": ["信息披露"],
            "action": "应当",
            "clause_text": "银行业金融机构应当按规定披露信息",
        },
    ]

    inner_clauses = [
        {
            "uid": "inner_1",
            "subjects": ["商业银行"],
            "objects": ["信用贷款"],
            "biz_dims": ["信贷管理"],
            "action": "禁止",
            "clause_text": "本行禁止向关系人发放信用贷款",
        },
        {
            "uid": "inner_2",
            "subjects": ["本行"],
            "objects": ["财务报告"],
            "biz_dims": ["信息披露"],
            "action": "披露",
            "clause_text": "本行每季度披露财务报告",
        },
    ]

    # 筛选配置（使用单个 config 参数）
    filter_config = {
        "hard_filter": {
            "biz_dim_match": ">= 1",
        },
        "soft_filter": {
            "min_score": 0.3,
            "score_weights": {
                "biz_dim_match": 0.30,
                "subject_match": 0.25,
                "object_match": 0.25,
                "action_match": 0.10,
                "text_similarity": 0.10,
            },
        },
    }

    # 执行筛选
    candidates = filter_candidates(
        outer_clauses,
        inner_clauses,
        filter_config,
    )

    # 验证筛选结果
    assert len(candidates) > 0

    # 验证每个候选对都有必要字段
    for candidate in candidates:
        assert "outer_clause" in candidate
        assert "inner_clause" in candidate
        assert "match_details" in candidate
        assert "similarity_score" in candidate
